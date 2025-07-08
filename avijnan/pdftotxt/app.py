#!/usr/bin/env python3
"""
Multilingual PDF-to-DOCX Converter (No GUI)
A robust converter that extracts text from PDFs and creates properly formatted DOCX files
with support for complex scripts and multilingual content.

Author: AI Software Engineer
Version: 1.0.0
License: MIT
"""

import os
import sys
import logging
import tempfile
import shutil
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass
from enum import Enum
import json
import re

# Core libraries
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import cv2
import numpy as np
from pdf2image import convert_from_path
import unicodedata
from langdetect import detect
from langdetect.lang_detect_exception import LangDetectException as LangDetectError

# DOCX creation
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

class TextExtractionMode(Enum):
    """Text extraction modes"""
    UNICODE_ONLY = "unicode_only"
    OCR_ONLY = "ocr_only"
    HYBRID = "hybrid"
    AUTO = "auto"

@dataclass
class ExtractionResult:
    """Result of text extraction"""
    text: str
    mode_used: TextExtractionMode
    confidence: float
    language_detected: Optional[str] = None
    pages_processed: int = 0
    unicode_pages: int = 0
    ocr_pages: int = 0
    warnings: List[str] = None
    formatting_info: Dict = None
    
    def __post_init__(self):
        if self.warnings is None:
            self.warnings = []
        if self.formatting_info is None:
            self.formatting_info = {}

class LanguageConfig:
    """Language configuration for OCR and text processing"""
    
    # Tesseract language codes mapping
    TESSERACT_LANGS = {
        'eng': 'English',
        'ben': 'Bengali',
        'hin': 'Hindi',
        'ara': 'Arabic',
        'jpn': 'Japanese',
        'chi_sim': 'Chinese (Simplified)',
        'chi_tra': 'Chinese (Traditional)',
        'tha': 'Thai',
        'kor': 'Korean',
        'tam': 'Tamil',
        'tel': 'Telugu',
        'kan': 'Kannada',
        'mal': 'Malayalam',
        'guj': 'Gujarati',
        'pan': 'Punjabi',
        'urd': 'Urdu',
        'nep': 'Nepali',
        'sin': 'Sinhala',
        'mya': 'Myanmar',
        'khm': 'Khmer',
        'lao': 'Lao',
        'vie': 'Vietnamese',
        'fas': 'Persian',
        'heb': 'Hebrew',
        'tur': 'Turkish',
        'rus': 'Russian',
        'deu': 'German',
        'fra': 'French',
        'spa': 'Spanish',
        'por': 'Portuguese',
        'ita': 'Italian',
        'nld': 'Dutch',
        'swe': 'Swedish',
        'nor': 'Norwegian',
        'dan': 'Danish',
        'fin': 'Finnish',
        'pol': 'Polish',
        'ces': 'Czech',
        'hun': 'Hungarian',
        'ron': 'Romanian',
        'ell': 'Greek',
        'bul': 'Bulgarian',
        'ukr': 'Ukrainian',
        'hrv': 'Croatian',
        'srp': 'Serbian',
        'slk': 'Slovak',
        'slv': 'Slovenian',
        'lit': 'Lithuanian',
        'lav': 'Latvian',
        'est': 'Estonian',
        'cat': 'Catalan',
        'eus': 'Basque',
        'glg': 'Galician',
        'mlt': 'Maltese',
        'gle': 'Irish',
        'cym': 'Welsh',
        'gla': 'Scottish Gaelic',
        'afr': 'Afrikaans',
        'swa': 'Swahili',
        'amh': 'Amharic',
        'msa': 'Malay',
        'ind': 'Indonesian',
        'tgl': 'Tagalog'
    }
    
    # Right-to-left languages
    RTL_LANGUAGES = {'ara', 'heb', 'fas', 'urd'}
    
    # Complex script languages
    COMPLEX_SCRIPTS = {
        'ben', 'hin', 'ara', 'jpn', 'chi_sim', 'chi_tra', 'tha', 'kor',
        'tam', 'tel', 'kan', 'mal', 'guj', 'pan', 'urd', 'nep', 'sin',
        'mya', 'khm', 'lao', 'heb', 'fas'
    }
    
    @classmethod
    def get_available_languages(cls) -> List[str]:
        """Get list of available Tesseract languages"""
        try:
            langs = pytesseract.get_languages(config='')
            return [lang for lang in langs if lang in cls.TESSERACT_LANGS]
        except Exception:
            return ['eng']
    
    @classmethod
    def detect_language_from_text(cls, text: str) -> Optional[str]:
        """Detect language from extracted text"""
        try:
            detected = detect(text)
            lang_mapping = {
                'en': 'eng', 'bn': 'ben', 'hi': 'hin', 'ar': 'ara',
                'ja': 'jpn', 'zh-cn': 'chi_sim', 'zh-tw': 'chi_tra',
                'th': 'tha', 'ko': 'kor', 'ta': 'tam', 'te': 'tel',
                'kn': 'kan', 'ml': 'mal', 'gu': 'guj', 'pa': 'pan',
                'ur': 'urd', 'ne': 'nep', 'si': 'sin', 'my': 'mya',
                'km': 'khm', 'lo': 'lao', 'vi': 'vie', 'fa': 'fas',
                'he': 'heb', 'tr': 'tur', 'ru': 'rus', 'de': 'deu',
                'fr': 'fra', 'es': 'spa', 'pt': 'por', 'it': 'ita',
                'nl': 'nld', 'sv': 'swe', 'no': 'nor', 'da': 'dan',
                'fi': 'fin', 'pl': 'pol', 'cs': 'ces', 'hu': 'hun',
                'ro': 'ron', 'el': 'ell', 'bg': 'bul', 'uk': 'ukr',
                'hr': 'hrv', 'sr': 'srp', 'sk': 'slk', 'sl': 'slv',
                'lt': 'lit', 'lv': 'lav', 'et': 'est', 'ca': 'cat',
                'eu': 'eus', 'gl': 'glg', 'mt': 'mlt', 'ga': 'gle',
                'cy': 'cym', 'gd': 'gla', 'af': 'afr', 'sw': 'swa',
                'am': 'amh', 'ms': 'msa', 'id': 'ind', 'tl': 'tgl'
            }
            return lang_mapping.get(detected, 'eng')
        except (LangDetectError, Exception):
            return None

class PDFTextExtractor:
    """Main PDF text extraction class"""
    
    def __init__(self, logger: logging.Logger = None):
        self.logger = logger or logging.getLogger(__name__)
        self.temp_dir = None
        self._verify_dependencies()
        
    def _verify_dependencies(self):
        """Verify all required dependencies are available"""
        missing_deps = []
        
        try:
            import fitz
        except ImportError:
            missing_deps.append("PyMuPDF")
        
        try:
            pytesseract.get_tesseract_version()
        except Exception:
            missing_deps.append("Tesseract OCR")
        
        try:
            from pdf2image import convert_from_path
        except ImportError:
            missing_deps.append("pdf2image")
        
        try:
            import cv2
        except ImportError:
            missing_deps.append("opencv-python")
        
        try:
            from docx import Document
        except ImportError:
            missing_deps.append("python-docx")
        
        if missing_deps:
            error_msg = f"Missing dependencies: {', '.join(missing_deps)}"
            self.logger.error(error_msg)
            raise ImportError(error_msg)
    
    def extract_text(self, pdf_path: str, mode: TextExtractionMode = TextExtractionMode.AUTO,
                    languages: List[str] = None, confidence_threshold: float = 0.6) -> ExtractionResult:
        """Extract text from PDF with formatting information"""
        self.logger.info(f"Starting text extraction from: {pdf_path}")
        
        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"PDF file not found: {pdf_path}")
        
        self.temp_dir = tempfile.mkdtemp(prefix="pdf_converter_")
        
        try:
            doc = fitz.open(pdf_path)
            total_pages = len(doc)
            
            if mode == TextExtractionMode.AUTO:
                mode = self._determine_extraction_mode(doc)
            
            if mode == TextExtractionMode.UNICODE_ONLY:
                result = self._extract_unicode_text_with_formatting(doc)
            elif mode == TextExtractionMode.OCR_ONLY:
                result = self._extract_ocr_text(doc, languages, confidence_threshold)
            elif mode == TextExtractionMode.HYBRID:
                result = self._extract_hybrid_text(doc, languages, confidence_threshold)
            else:
                raise ValueError(f"Unsupported extraction mode: {mode}")
            
            if result.text.strip() and not result.language_detected:
                result.language_detected = LanguageConfig.detect_language_from_text(result.text)
            
            result.text = self._post_process_text(result.text, result.language_detected)
            
            doc.close()
            return result
            
        except Exception as e:
            self.logger.error(f"Error during text extraction: {str(e)}")
            raise
        finally:
            if self.temp_dir and os.path.exists(self.temp_dir):
                shutil.rmtree(self.temp_dir)
    
    def _determine_extraction_mode(self, doc: fitz.Document) -> TextExtractionMode:
        """Determine the best extraction mode"""
        total_pages = len(doc)
        sample_pages = min(3, total_pages)
        
        unicode_text_found = 0
        total_text_length = 0
        
        for page_num in range(sample_pages):
            page = doc[page_num]
            text = page.get_text()
            
            if text.strip():
                unicode_text_found += 1
                total_text_length += len(text)
        
        if unicode_text_found >= sample_pages * 0.7 and total_text_length > 100:
            return TextExtractionMode.UNICODE_ONLY
        elif unicode_text_found == 0:
            return TextExtractionMode.OCR_ONLY
        else:
            return TextExtractionMode.HYBRID
    
    def _extract_unicode_text_with_formatting(self, doc: fitz.Document) -> ExtractionResult:
        """Extract Unicode text with formatting information"""
        extracted_text = []
        pages_processed = 0
        unicode_pages = 0
        warnings = []
        formatting_info = {
            'paragraphs': [],
            'fonts': {},
            'styles': []
        }
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            # Extract text blocks with formatting
            blocks = page.get_text("dict")
            page_text = []
            
            for block in blocks.get("blocks", []):
                if "lines" in block:
                    block_text = []
                    for line in block["lines"]:
                        line_text = []
                        for span in line.get("spans", []):
                            text = span.get("text", "")
                            if text.strip():
                                line_text.append(text)
                                
                                # Store font information
                                font_info = {
                                    'font': span.get('font', ''),
                                    'size': span.get('size', 12),
                                    'bold': 'Bold' in span.get('font', ''),
                                    'italic': 'Italic' in span.get('font', ''),
                                }
                                font_key = f"{font_info['font']}_{font_info['size']}"
                                formatting_info['fonts'][font_key] = font_info
                        
                        if line_text:
                            block_text.append(" ".join(line_text))
                    
                    if block_text:
                        paragraph = "\n".join(block_text)
                        page_text.append(paragraph)
                        formatting_info['paragraphs'].append({
                            'text': paragraph,
                            'page': page_num + 1
                        })
            
            if page_text:
                page_content = "\n\n".join(page_text)
                if self._is_text_quality_good(page_content):
                    extracted_text.append(page_content)
                    unicode_pages += 1
                else:
                    warnings.append(f"Page {page_num + 1}: Low quality Unicode text")
            else:
                warnings.append(f"Page {page_num + 1}: No Unicode text found")
            
            pages_processed += 1
        
        full_text = "\n\n".join(extracted_text)
        confidence = unicode_pages / pages_processed if pages_processed > 0 else 0
        
        return ExtractionResult(
            text=full_text,
            mode_used=TextExtractionMode.UNICODE_ONLY,
            confidence=confidence,
            pages_processed=pages_processed,
            unicode_pages=unicode_pages,
            ocr_pages=0,
            warnings=warnings,
            formatting_info=formatting_info
        )
    
    def _extract_ocr_text(self, doc: fitz.Document, languages: List[str] = None, 
                         confidence_threshold: float = 0.6) -> ExtractionResult:
        """Extract text using OCR"""
        if not languages:
            languages = ['eng']
        
        images = convert_from_path(doc.name, dpi=300, thread_count=4)
        
        extracted_text = []
        pages_processed = 0
        ocr_pages = 0
        warnings = []
        total_confidence = 0
        
        lang_string = '+'.join(languages)
        config = f'--oem 3 --psm 6 -l {lang_string}'
        is_rtl = any(lang in LanguageConfig.RTL_LANGUAGES for lang in languages)
        
        for page_num, image in enumerate(images):
            try:
                cv_image = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
                processed_image = self._preprocess_image_for_ocr(cv_image)
                
                ocr_data = pytesseract.image_to_data(processed_image, config=config, 
                                                   output_type=pytesseract.Output.DICT)
                
                page_text, page_confidence = self._extract_text_with_confidence(
                    ocr_data, confidence_threshold, is_rtl
                )
                
                if page_text.strip():
                    extracted_text.append(page_text)
                    ocr_pages += 1
                    total_confidence += page_confidence
                else:
                    warnings.append(f"Page {page_num + 1}: No text extracted via OCR")
                
                pages_processed += 1
                
            except Exception as e:
                warnings.append(f"Page {page_num + 1}: OCR error - {str(e)}")
                pages_processed += 1
        
        full_text = "\n\n".join(extracted_text)
        overall_confidence = total_confidence / ocr_pages if ocr_pages > 0 else 0
        
        return ExtractionResult(
            text=full_text,
            mode_used=TextExtractionMode.OCR_ONLY,
            confidence=overall_confidence,
            pages_processed=pages_processed,
            unicode_pages=0,
            ocr_pages=ocr_pages,
            warnings=warnings
        )
    
    def _extract_hybrid_text(self, doc: fitz.Document, languages: List[str] = None,
                           confidence_threshold: float = 0.6) -> ExtractionResult:
        """Extract text using hybrid approach"""
        if not languages:
            languages = ['eng']
        
        extracted_text = []
        pages_processed = 0
        unicode_pages = 0
        ocr_pages = 0
        warnings = []
        total_confidence = 0
        
        images = convert_from_path(doc.name, dpi=300, thread_count=4)
        
        lang_string = '+'.join(languages)
        config = f'--oem 3 --psm 6 -l {lang_string}'
        is_rtl = any(lang in LanguageConfig.RTL_LANGUAGES for lang in languages)
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            unicode_text = page.get_text()
            
            if unicode_text.strip() and self._is_text_quality_good(unicode_text):
                extracted_text.append(unicode_text)
                unicode_pages += 1
                total_confidence += 0.95
            else:
                try:
                    image = images[page_num]
                    cv_image = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
                    processed_image = self._preprocess_image_for_ocr(cv_image)
                    
                    ocr_data = pytesseract.image_to_data(processed_image, config=config,
                                                       output_type=pytesseract.Output.DICT)
                    
                    ocr_text, ocr_confidence = self._extract_text_with_confidence(
                        ocr_data, confidence_threshold, is_rtl
                    )
                    
                    if ocr_text.strip():
                        extracted_text.append(ocr_text)
                        ocr_pages += 1
                        total_confidence += ocr_confidence
                    else:
                        warnings.append(f"Page {page_num + 1}: No text extracted")
                        
                except Exception as e:
                    warnings.append(f"Page {page_num + 1}: OCR error - {str(e)}")
            
            pages_processed += 1
        
        full_text = "\n\n".join(extracted_text)
        overall_confidence = total_confidence / pages_processed if pages_processed > 0 else 0
        
        return ExtractionResult(
            text=full_text,
            mode_used=TextExtractionMode.HYBRID,
            confidence=overall_confidence,
            pages_processed=pages_processed,
            unicode_pages=unicode_pages,
            ocr_pages=ocr_pages,
            warnings=warnings
        )
    
    def _preprocess_image_for_ocr(self, image: np.ndarray) -> np.ndarray:
        """Preprocess image for better OCR"""
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        denoised = cv2.fastNlMeansDenoising(gray)
        thresh = cv2.adaptiveThreshold(denoised, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                     cv2.THRESH_BINARY, 11, 2)
        kernel = np.ones((1, 1), np.uint8)
        cleaned = cv2.morphologyEx(thresh, cv2.MORPH_CLOSE, kernel)
        return cleaned
    
    def _extract_text_with_confidence(self, ocr_data: Dict, confidence_threshold: float,
                                    is_rtl: bool = False) -> Tuple[str, float]:
        """Extract text with confidence filtering"""
        words = []
        confidences = []
        
        for i, conf in enumerate(ocr_data['conf']):
            if int(conf) > confidence_threshold * 100:
                text = ocr_data['text'][i].strip()
                if text:
                    words.append(text)
                    confidences.append(int(conf) / 100.0)
        
        if not words:
            return "", 0.0
        
        text = " ".join(words)
        avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0
        
        return text, avg_confidence
    
    def _is_text_quality_good(self, text: str) -> bool:
        """Check text quality"""
        if not text.strip() or len(text) < 10:
            return False
        
        letters = sum(1 for c in text if c.isalpha())
        letter_ratio = letters / len(text)
        return letter_ratio > 0.3
    
    def _post_process_text(self, text: str, language: str = None) -> str:
        """Post-process text"""
        if not text.strip():
            return text
        
        text = unicodedata.normalize('NFC', text)
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        text = re.sub(r'\t+', ' ', text)
        
        if language and language in LanguageConfig.RTL_LANGUAGES:
            text = '\u202B' + text + '\u202C'
        
        return text.strip()

class DOCXCreator:
    """Create DOCX files from extracted text"""
    
    def __init__(self, logger: logging.Logger = None):
        self.logger = logger or logging.getLogger(__name__)
    
    def create_docx(self, extraction_result: ExtractionResult, output_path: str, 
                   preserve_formatting: bool = True):
        """Create DOCX file from extraction result"""
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = f"Converted from PDF"
        doc.core_properties.author = "PDF-to-DOCX Converter"
        
        # Configure styles for multilingual support
        self._setup_multilingual_styles(doc, extraction_result.language_detected)
        
        if preserve_formatting and extraction_result.formatting_info:
            self._add_formatted_content(doc, extraction_result)
        else:
            self._add_simple_content(doc, extraction_result)
        
        # Save document
        doc.save(output_path)
        self.logger.info(f"DOCX file saved to: {output_path}")
    
    def _setup_multilingual_styles(self, doc: Document, language: str = None):
        """Setup styles for multilingual content"""
        styles = doc.styles
        
        # Create multilingual paragraph style
        try:
            style = styles.add_style('Multilingual', WD_STYLE_TYPE.PARAGRAPH)
            font = style.font
            font.name = 'Arial Unicode MS'
            font.size = Pt(11)
            
            # Set paragraph alignment based on language
            if language and language in LanguageConfig.RTL_LANGUAGES:
                style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
        except Exception as e:
            self.logger.warning(f"Could not create multilingual style: {e}")
    
    def _add_formatted_content(self, doc: Document, extraction_result: ExtractionResult):
        """Add content with formatting preservation"""
        formatting_info = extraction_result.formatting_info
        
        # Add title
        title = doc.add_heading('Converted PDF Content', level=1)
        
        # Add paragraphs with formatting
        if formatting_info and 'paragraphs' in formatting_info:
            for para_info in formatting_info['paragraphs']:
                para_text = para_info['text'].strip()
                if para_text:
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(para_text)
                    
                    # Apply font formatting if available
                    if formatting_info.get('fonts'):
                        # Use first available font info
                        font_info = next(iter(formatting_info['fonts'].values()))
                        if font_info.get('bold'):
                            run.bold = True
                        if font_info.get('italic'):
                            run.italic = True
                    
                    # Apply multilingual style
                    try:
                        paragraph.style = 'Multilingual'
                    except:
                        pass
        else:
            # Fallback to simple content
            self._add_simple_content(doc, extraction_result)
    
    def _add_simple_content(self, doc: Document, extraction_result: ExtractionResult):
        """Add simple content without formatting"""
        # Add title
        title = doc.add_heading('Converted PDF Content', level=1)
        
        # Split text into paragraphs
        text = extraction_result.text
        paragraphs = text.split('\n\n')
        
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                paragraph = doc.add_paragraph(para_text)
                
                # Apply multilingual style
                try:
                    paragraph.style = 'Multilingual'
                except:
                    pass

class PDFToDOCXConverter:
    """Main converter class"""
    
    def __init__(self, silent: bool = False):
        self.setup_logging(silent)
        self.extractor = PDFTextExtractor(self.logger)
        self.docx_creator = DOCXCreator(self.logger)
    
    def setup_logging(self, silent: bool = False):
        """Setup logging"""
        if silent:
            logging.basicConfig(level=logging.ERROR)
        else:
            logging.basicConfig(
                level=logging.INFO,
                format='%(asctime)s - %(levelname)s - %(message)s'
            )
        self.logger = logging.getLogger(__name__)
    
    def convert(self, pdf_path: str, output_path: str = None, 
               mode: TextExtractionMode = TextExtractionMode.AUTO,
               languages: List[str] = None, confidence_threshold: float = 0.6,
               preserve_formatting: bool = True) -> str:
        """Convert PDF to DOCX and return the output path"""
        
        if not output_path:
            output_path = Path(pdf_path).stem + '.docx'
        
        self.logger.info(f"Converting PDF to DOCX: {pdf_path} -> {output_path}")
        
        try:
            # Extract text
            extraction_result = self.extractor.extract_text(
                pdf_path, mode, languages, confidence_threshold
            )
            
            # Create DOCX
            self.docx_creator.create_docx(
                extraction_result, output_path, preserve_formatting
            )
            
            self.logger.info("Conversion completed successfully")
            return os.path.abspath(output_path)
            
        except Exception as e:
            self.logger.error(f"Conversion failed: {str(e)}")
            raise

def convert_pdf_to_docx(pdf_path: str, output_path: str = None, 
                       mode: str = "auto", languages: List[str] = None,
                       confidence_threshold: float = 0.6,
                       preserve_formatting: bool = True,
                       silent: bool = False) -> str:
    """
    Simple function to convert PDF to DOCX and return the output path.
    
    Args:
        pdf_path: Path to the PDF file
        output_path: Path for the output DOCX file (optional)
        mode: Extraction mode ('auto', 'unicode_only', 'ocr_only', 'hybrid')
        languages: List of language codes for OCR (default: ['eng'])
        confidence_threshold: OCR confidence threshold (default: 0.6)
        preserve_formatting: Whether to preserve formatting (default: True)
        silent: Whether to suppress logging output (default: False)
    
    Returns:
        str: Absolute path to the created DOCX file
    """
    converter = PDFToDOCXConverter(silent=silent)
    
    if not languages:
        languages = ['eng']
    
    return converter.convert(
        pdf_path=pdf_path,
        output_path=output_path,
        mode=TextExtractionMode(mode),
        languages=languages,
        confidence_threshold=confidence_threshold,
        preserve_formatting=preserve_formatting
    )

